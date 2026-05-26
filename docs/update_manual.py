"""更新用户手册 — 添加 MCP Server 配置章节并更新版本信息

用法: python docs/update_manual.py
会修改 docs/ShineHeKnowledge_UserManual_v1.0.0.docx
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from pathlib import Path
from docx import Document
from lxml import etree

DOC_PATH = Path(__file__).parent / "ShineHeKnowledge_UserManual_v1.0.0.docx"

# ---- MCP 配置 JSON 模板 ----

CLAUDE_CODE_CONFIG = (
    '// 在 ~/.claude.json 的 "mcpServers" 中添加：\n'
    '"shinehe-kb": {\n'
    '  "command": "python",\n'
    '  "args": ["<项目路径>/run_mcp.py"]\n'
    '}'
)

CURSOR_CONFIG = (
    '// 在 Cursor Settings → MCP 中添加：\n'
    '{\n'
    '  "mcpServers": {\n'
    '    "shinehe-kb": {\n'
    '      "command": "python",\n'
    '      "args": ["<项目路径>/run_mcp.py"]\n'
    '    }\n'
    '  }\n'
    '}'
)

MCP_STDIO_CONFIG = (
    '// stdio 模式（推荐，AI 工具自动管理生命周期）：\n'
    '{\n'
    '  "shinehe-kb": {\n'
    '    "command": "python",\n'
    '    "args": ["<项目路径>/run_mcp.py"],\n'
    '    "transport": "stdio"\n'
    '  }\n'
    '}'
)

MCP_HTTP_CONFIG = (
    '// HTTP 模式（需先手动启动：python run_mcp.py -t streamable-http）：\n'
    '{\n'
    '  "shinehe-kb": {\n'
    '    "url": "http://127.0.0.1:9000/mcp",\n'
    '    "transport": "streamable-http"\n'
    '  }\n'
    '}'
)


def _insert_paragraphs_before(doc, target_elem, sections):
    """在 target_elem 前按正序插入段落列表"""
    body = doc.element.body
    for style_name, text in sections:
        new_p = doc.add_paragraph("", style=style_name)
        if text:
            run = new_p.add_run(text)
            if style_name == "Normal" and text.startswith(("//", "{")):
                run.font.name = "Consolas"
                run.font.size = None  # 跟随段落
        body.remove(new_p._element)
        target_elem.addprevious(new_p._element)


def add_mcp_section(doc):
    """在「5. 配置说明」之前插入 MCP Server 章节"""

    target_elem = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and "5. 配置说明" in p.text:
            target_elem = p._element
            break

    if target_elem is None:
        print("ERROR: 未找到「5. 配置说明」章节")
        return

    sections = [
        ("Heading 1", "4.6 MCP Server 集成"),
        ("Normal", (
            "ShineHeKnowledge 内置 MCP（Model Context Protocol）服务器，可被 Claude Code、"
            "TeleClaw、Cursor 等 AI 工具直接调用，实现知识库的智能检索与问答。"
        )),
        ("Normal", ""),
        ("Heading 2", "4.6.1 启动 MCP Server"),
        ("Normal", "MCP Server 支持三种传输模式："),
        ("List Bullet", "stdio 模式（默认）：适合被 AI 工具直接启动，无需额外运行服务"),
        ("List Bullet", "streamable-http 模式：独立启动 HTTP 服务，适合多客户端连接"),
        ("List Bullet", "sse 模式：SSE 长连接模式，适合 Web 客户端"),
        ("Normal", ""),
        ("Normal", "启动命令："),
        ("Normal", "python run_mcp.py                                    # stdio 模式（默认）"),
        ("Normal", "python run_mcp.py -t streamable-http -p 9000         # HTTP 模式"),
        ("Normal", "python run_mcp.py -t sse -p 9000                     # SSE 模式"),
        ("Normal", ""),
        ("Heading 2", "4.6.2 可用工具列表"),
        ("Normal", "连接 MCP Server 后，AI 工具可调用以下工具："),
        ("List Bullet", "search — 语义向量搜索，查找含义最相关的知识内容"),
        ("List Bullet", "search_fulltext — FTS5 全文关键词搜索，精确匹配"),
        ("List Bullet", "ask — RAG 智能问答，自动检索 + LLM 生成回答"),
        ("List Bullet", "create — 创建知识条目（自动分块、向量化索引）"),
        ("List Bullet", "read — 读取指定知识条目的完整信息"),
        ("List Bullet", "update — 更新知识条目（标题、内容、标签）"),
        ("List Bullet", "delete — 删除知识条目及其向量索引"),
        ("List Bullet", "list_knowledge — 列出知识条目（支持筛选、分页、排序）"),
        ("List Bullet", "tags — 获取所有标签列表"),
        ("List Bullet", "ingest_file — 解析本地文件并导入知识库（支持 PDF、DOCX、PPT、TXT、MD 等）"),
        ("Normal", ""),
        ("Heading 2", "4.6.3 客户端配置"),
        ("Normal", ""),
        ("Normal", "【Claude Code】在 ~/.claude.json 的 mcpServers 中添加："),
        ("Normal", CLAUDE_CODE_CONFIG),
        ("Normal", ""),
        ("Normal", "【Cursor】在 Cursor Settings → MCP 中添加："),
        ("Normal", CURSOR_CONFIG),
        ("Normal", ""),
        ("Normal", "【TeleClaw 等其他 MCP 客户端】stdio 模式配置："),
        ("Normal", MCP_STDIO_CONFIG),
        ("Normal", ""),
        ("Normal", "如需 HTTP 模式，先手动启动服务（python run_mcp.py -t streamable-http），然后配置："),
        ("Normal", MCP_HTTP_CONFIG),
        ("Normal", ""),
        ("Normal", (
            "注意：请将 <项目路径> 替换为 run_mcp.py 的实际完整路径。"
            "如果 python 命令不在 PATH 中，请使用 Python 解释器的完整路径，"
            "如 C:/Users/xxx/AppData/Local/Programs/Python/Python314/python.exe。"
        )),
        ("Normal", ""),
    ]

    # 用 addprevious 逐个插入 — 不需要反序，addprevious 每次插到 target 前
    body = doc.element.body
    for style_name, text in sections:
        new_p = doc.add_paragraph("", style=style_name)
        if text:
            run = new_p.add_run(text)
            if style_name == "Normal" and text.startswith(("//", "{")):
                run.font.name = "Consolas"
        body.remove(new_p._element)
        target_elem.addprevious(new_p._element)


def update_toc(doc):
    """更新目录"""
    for i, p in enumerate(doc.paragraphs):
        if "4.5 插件扩展" in p.text and "Heading" not in p.style.name:
            # 找到下一行作为锚点
            anchor = None
            for j in range(i + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():
                    anchor = doc.paragraphs[j]._element
                    break
            if anchor is None:
                break
            toc_entry = doc.add_paragraph("", style="Normal")
            toc_entry.add_run("   4.6 MCP Server 集成")
            body = doc.element.body
            body.remove(toc_entry._element)
            anchor.addprevious(toc_entry._element)
            break


def update_version(doc):
    """更新版本日志"""
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and "v1.0.0 (" in p.text and "v1.0.0+" not in p.text:
            target = p._element
            body = doc.element.body

            entries = [
                ("Heading 2", "v1.0.0+ (2026-05) — 持续更新"),
                ("List Bullet", "UI 全面重构为 Notion/Obsidian 现代简约风格"),
                ("List Bullet", "侧边栏改为 200px 宽文字导航，统一品牌标识「世恒的图书馆」"),
                ("List Bullet", "MCP Server 添加心跳机制，GUI 侧边栏实时显示 MCP 连接状态"),
                ("List Bullet", "新增 PPT/PPTX 文件导入支持"),
                ("List Bullet", "LLM 调用状态指示灯（侧边栏底部）"),
                ("List Bullet", "启动性能优化（chromadb 等重型库延迟加载）"),
                ("List Bullet", "知识分类体系与目录浏览功能"),
                ("List Bullet", "用户手册添加 MCP Server 集成配置说明"),
                ("Normal", ""),
            ]

            for style_name, text in entries:
                new_p = doc.add_paragraph("", style=style_name)
                new_p.add_run(text)
                body.remove(new_p._element)
                target.addprevious(new_p._element)
            break


def update_cover(doc):
    """更新封面版本号和日期"""
    for p in doc.paragraphs:
        if "Version 1.0.0" in p.text:
            for run in p.runs:
                if "1.0.0" in run.text and "+" not in run.text:
                    run.text = run.text.replace("1.0.0", "1.0.0+")
        if "2026 年 04 月" in p.text:
            for run in p.runs:
                if "04 月" in run.text:
                    run.text = run.text.replace("04 月", "05 月")


def update_features(doc):
    """更新核心特性：替换旧描述"""
    for p in doc.paragraphs:
        if p.style.name == "List Bullet" and "暗色科幻界面" in p.text:
            for run in p.runs:
                if "暗色科幻界面" in run.text:
                    run.text = "MCP Server 集成：支持 Claude Code、TeleClaw、Cursor 等 AI 工具直接调用"
            break


def main():
    print(f"正在更新文档: {DOC_PATH}")
    doc = Document(str(DOC_PATH))

    update_cover(doc)
    update_features(doc)
    update_toc(doc)
    add_mcp_section(doc)
    update_version(doc)

    doc.save(str(DOC_PATH))
    print("文档更新完成！")


if __name__ == "__main__":
    main()
