"""生成 ShineHeKnowledge 用户说明文档 (DOCX)

用法: python scripts/build_docs.py
输出: docs/ShineHeKnowledge_UserManual_vX.X.X.docx
"""
import sys
import os
from datetime import datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
from src.version import APP_NAME, VERSION

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _add_table(doc, headers, rows):
    """快速生成 Table Grid 表格。"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    return table


def _consolas(doc, text):
    """添加 Consolas 字体段落。"""
    p = doc.add_paragraph(text)
    p.runs[0].font.name = "Consolas"
    return p


def build_doc():
    doc = Document()

    # ---- 样式设置 ----
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Microsoft YaHei"
        hs.font.color.rgb = RGBColor(0, 100, 60)

    # ---- 封面 ----
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(APP_NAME)
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 140, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version {VERSION}")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("用户说明文档")
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%Y 年 %m 月"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # ---- 目录页 ----
    doc.add_heading("目录", level=1)
    toc_items = [
        "1. 系统简介",
        "2. 安装指南",
        "   2.1 Windows 安装包部署",
        "   2.2 Docker 容器部署",
        "   2.3 从源码运行",
        "3. 快速上手",
        "4. 功能说明",
        "   4.1 知识库管理",
        "   4.2 智能搜索",
        "   4.3 RAG 智能问答（6 阶段管线）",
        "   4.4 Wiki 知识沉淀",
        "   4.5 知识图谱与 Block 引用",
        "   4.6 DSL 结构化查询",
        "   4.7 MCP Server 集成",
        "   4.8 REST API 接口",
        "   4.9 安全与审计",
        "5. 配置说明",
        "   5.1 LLM 配置",
        "   5.2 Embedding 配置",
        "   5.3 RAG 管线参数",
        "6. API 接口参考",
        "7. MCP 工具一览",
        "8. 版本更新日志",
        "9. 常见问题",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ==== 1. 系统简介 ====
    doc.add_heading("1. 系统简介", level=1)
    doc.add_paragraph(
        f"{APP_NAME} 是一款本地部署的多模态知识库管理系统，支持文档管理、RAG 智能问答、"
        "Wiki 知识沉淀、知识图谱可视化、DSL 结构化查询等功能。"
        "系统提供桌面 GUI（PySide6）、REST API（FastAPI）和 MCP Server（FastMCP）三种访问方式，"
        "适配不同使用场景。所有数据本地存储，API Key 通过系统 Keychain 安全保管，JWT 认证，"
        "MCP 协议集成，支持 Claude Code / Cursor / Cline 等主流 AI Agent 一键接入。"
    )
    doc.add_heading("核心特性", level=2)
    features = [
        "多格式文档导入：PDF、DOCX、XLSX、CSV、PPTX、HTML、TXT、Markdown、代码、图片等",
        "混合搜索（向量 + 关键词 + RRF 融合）：语义搜索与全文检索结合，支持 Reranker 重排序",
        "6 阶段可配置 RAG 管线：查询改写 → Wiki 检索 → 向量搜索 → 重排序 → 生成 → 后处理",
        "Wiki 知识沉淀：原始文档自动编译为结构化 Wiki 页面，支持审核工作流和自动分类",
        "知识图谱可视化：6 种关系类型、2D 力导向图可视化、知识目录浏览、LLM 自动分类（15 大类 57 小类）",
        "Block 级引用体系：支持段落级嵌入（Transclusion）、链接扩展、DSL 查询语言",
        "异步任务系统：大文件自动转异步导入，支持任务状态追踪和取消",
        "安全与审计：操作日志、预览模式、撤销操作、软删除与恢复",
        "MCP Server：51 个原始工具 + 51 个命名空间别名 + 3 个资源 + 5 个 Prompt，支持 stdio / streamable-http / sse",
        "跨平台：Python 全栈架构，Windows 安装包、Docker 容器、源码运行三种部署方式",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    # ==== 2. 安装指南 ====
    doc.add_heading("2. 安装指南", level=1)

    doc.add_heading("2.1 Windows 安装包部署", level=2)
    doc.add_paragraph("适合不熟悉命令行的用户。")
    steps = [
        "双击 ShineHeKnowledge_vX.X.X_Setup.exe 运行安装向导",
        "选择安装目录，点击「下一步」完成安装",
        "安装完成后，桌面会出现 ShineHeKnowledge 快捷方式",
        "双击快捷方式启动应用",
        "首次使用请点击左下角「设置」按钮配置 API Key",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_heading("2.2 Docker 容器部署", level=2)
    doc.add_paragraph("适合需要服务化部署或跨平台运行的用户。")
    doc.add_paragraph("确保已安装 Docker 和 Docker Compose，然后执行：")
    _consolas(doc, "docker-compose up -d shinehe-api")
    doc.add_paragraph("API 服务将在 http://localhost:8000 启动，访问 /docs 查看 API 文档。")

    doc.add_heading("2.3 从源码运行", level=2)
    doc.add_paragraph("需要 Python 3.10+ 环境。")
    for s in [
        "git clone <repo-url> && cd knowledge-base",
        'pip install -e .                # MCP 核心',
        'pip install -e ".[parsers]"    # 含文件解析（PDF/DOCX/图片等）',
        'pip install -e ".[all]"        # 含 GUI + API + 解析器 + Wiki',
        "",
        "python main.py                  # 桌面 GUI",
        "python run_api.py               # REST API（端口 8000）",
        "python run_mcp.py               # MCP Server（stdio 模式）",
    ]:
        if s:
            _consolas(doc, s)
        else:
            doc.add_paragraph()

    # ==== 3. 快速上手 ====
    doc.add_heading("3. 快速上手", level=1)

    doc.add_heading("3.1 配置 API", level=2)
    doc.add_paragraph(
        "启动应用后，点击左侧「设置」按钮（或在 API 中调用 /api/auth/register 注册账户）。"
        "在设置界面填写："
    )
    for c in [
        "供应商名称：如 deepseek、zhipu、moonshot 等",
        "API Key：从供应商获取的密钥",
        "API 地址：供应商的 OpenAI 兼容接口地址",
        "模型：如 deepseek-chat、glm-4-flash 等",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("3.2 导入知识", level=2)
    doc.add_paragraph(
        "在知识库界面点击「导入文件」按钮，选择一个或多个文件（支持 PDF、Word、Excel、Markdown 等），"
        "可选填标签，点击「开始导入」。系统将自动解析文件内容、分块、向量化并存储。"
        "大文件会自动转为异步任务，可在「任务管理」中查看进度。"
    )

    doc.add_heading("3.3 智能问答", level=2)
    doc.add_paragraph(
        "切换到「智能问答」页面，输入问题后按 Enter 发送。"
        "系统会自动检索最相关的知识内容，结合 LLM 生成回答，并标注引用来源和知识图谱。"
    )

    doc.add_heading("3.4 MCP 快速接入", level=2)
    doc.add_paragraph(
        "如果使用 Claude Code、Cursor 等 AI 编码工具，可通过 MCP 协议直接调用知识库："
    )
    _consolas(doc, "claude mcp add -t http shinehe-kb http://127.0.0.1:9000/mcp -s user")
    doc.add_paragraph(
        "更多配置模板见项目 mcp_config_templates/ 目录，支持 Claude Code / Cursor / Cline / "
        "Windsurf / Roo Code / OpenCode / TeleClaw 等主流 Agent。"
    )

    # ==== 4. 功能说明 ====
    doc.add_heading("4. 功能说明", level=1)

    # 4.1 知识库管理
    doc.add_heading("4.1 知识库管理", level=2)
    doc.add_paragraph("支持以下操作：")
    for m in [
        "创建：手动输入或从文件导入（支持 PDF/DOCX/XLSX/PPTX/HTML/TXT/MD/图片等）",
        "编辑：修改标题、内容、标签",
        "版本控制：每次编辑自动创建版本快照，可查看历史版本并一键恢复",
        "分类与标签：支持多标签，按标签筛选",
        "批量导出：通过 API 按标签或 ID 批量导出为 JSON",
        "软删除与恢复：删除后进入回收站，可随时恢复",
    ]:
        doc.add_paragraph(m, style="List Bullet")

    # 4.2 智能搜索
    doc.add_heading("4.2 智能搜索", level=2)
    doc.add_paragraph(
        "系统提供两种互补的搜索方式：\n"
        "• 语义向量搜索：通过 Embedding 模型将查询转换为向量，在向量空间中查找最相关的知识块\n"
        "• FTS5 全文搜索：基于 SQLite FTS5 引擎的关键词精确匹配，支持中英文分词\n\n"
        "两种搜索通过 RRF（Reciprocal Rank Fusion）融合，向量搜索权重 0.7，关键词权重 0.3，"
        "可配置。搜索结果可按更新时间、创建时间、标题排序。"
    )

    # 4.3 RAG 智能问答
    doc.add_heading("4.3 RAG 智能问答（6 阶段管线）", level=2)
    doc.add_paragraph(
        "RAG（Retrieval-Augmented Generation）采用可配置的 6 阶段管线架构："
    )
    pipeline_stages = [
        ("1. 查询改写（query_rewrite）", "LLM 生成多个查询变体，扩大召回范围"),
        ("2. Wiki 检索（wiki_retrieval）", "优先检索 Wiki 已综合验证的结构化知识"),
        ("3. 向量搜索（vector_search）", "混合搜索（向量 + 关键词 + RRF 融合），支持 Agentic Router 自动选择搜索策略"),
        ("4. 重排序（rerank）", "LLM 打分重排序，过滤低质量结果"),
        ("5. 生成（generate）", "LLM 基于检索上下文生成回答，支持流式输出"),
        ("6. 后处理（postprocess）", "去重、截断、格式化输出"),
    ]
    for name, desc in pipeline_stages:
        doc.add_paragraph(f"{name}：{desc}", style="List Bullet")
    doc.add_paragraph(
        "管线阶段可通过 config.yaml 的 rag.pipeline.stages 配置，支持自定义阶段扩展。"
        "Agentic Router 可自动识别结构化查询意图（如标签过滤、属性比较），走 DSL 捷径。"
    )

    # 4.4 Wiki 知识沉淀
    doc.add_heading("4.4 Wiki 知识沉淀", level=2)
    doc.add_paragraph(
        "Wiki 系统将原始文档自动编译为结构化 Wiki 页面，实现知识的持续沉淀和复利增长："
    )
    for w in [
        "自动编译：导入文件时自动触发 Wiki 编译（可通过 wiki.auto_compile 配置关闭）",
        "交叉引用发现：LLM 自动发现概念间的链接关系，构建知识网络",
        "合并逻辑：当新知识与已有 Wiki 页面相关时自动合并，避免重复",
        "问答回存：好的 RAG 问答结果可保存为 Wiki 页面",
        "审核工作流：draft → review → published → deprecated 四阶段状态机",
        "健康检查：wiki_lint 工具检测孤立页面、过时信息和损坏链接",
        "LLM 自动分类：支持 15 大类 57 小类的知识分类体系",
    ]:
        doc.add_paragraph(w, style="List Bullet")

    # 4.5 知识图谱与 Block 引用
    doc.add_heading("4.5 知识图谱与 Block 引用", level=2)
    doc.add_paragraph(
        "基于本地 Markdown 文件系统的知识图谱服务，支持："
    )
    for g in [
        "6 种关系类型：引用、关联、父子、兄弟、相似、顺序",
        "2D 力导向图可视化：交互式知识图谱浏览",
        "知识目录浏览：LLM 自动分类构建的树形目录",
        "Block 级引用：段落级嵌入（Transclusion），支持链接扩展和上下文窗口",
        "图谱遍历：从指定节点出发，多跳遍历知识网络",
        "数据导出：数据库知识条目可批量导出为 Markdown 图文件",
    ]:
        doc.add_paragraph(g, style="List Bullet")

    # 4.6 DSL 结构化查询
    doc.add_heading("4.6 DSL 结构化查询", level=2)
    doc.add_paragraph(
        "QuerySpec DSL 提供精确的知识库查询能力，支持三种执行模式："
    )
    for d in [
        "structured（结构化过滤）：按标签、属性条件组合过滤（and/or/not 逻辑，eq/ne/gt/lt/in/contains/like 操作符）",
        "graph（图遍历）：从指定节点出发，多跳遍历知识图谱",
        "hybrid（混合搜索）：语义向量 + 结构化过滤结合",
    ]:
        doc.add_paragraph(d, style="List Bullet")
    doc.add_paragraph(
        "可通过 MCP 工具 route_query 自动分析查询意图并推荐执行模式，"
        "也可通过 execute_query / ask_with_query 显式控制检索阶段。"
    )

    # 4.7 MCP Server 集成
    doc.add_heading("4.7 MCP Server 集成", level=2)
    doc.add_paragraph(
        "ShineHeKnowledge 内置 MCP（Model Context Protocol）服务器，可被 Claude Code、Cursor、"
        "Cline、Windsurf、Roo Code、OpenCode、TeleClaw 等 AI 工具直接调用。"
    )
    doc.add_heading("传输模式", level=3)
    for t in [
        "stdio 模式：适合被 AI 工具直接启动，无需额外运行服务",
        "streamable-http 模式（推荐）：独立启动 HTTP 服务，比 stdio 更稳定，不受管道断连影响",
        "sse 模式：SSE 长连接模式，适合 Web 客户端",
    ]:
        doc.add_paragraph(t, style="List Bullet")
    doc.add_heading("启动命令", level=3)
    for cmd in [
        "python run_mcp.py                                    # stdio 模式（默认）",
        "python run_mcp.py -t streamable-http -p 9000         # HTTP 模式",
        "python scripts/mcp_service.py start                  # 后台常驻（关闭终端不影响）",
    ]:
        _consolas(doc, cmd)
    doc.add_heading("客户端配置", level=3)
    doc.add_paragraph(
        "配置模板位于 mcp_config_templates/ 目录，覆盖主流 AI 编码工具。"
        "也可运行交互式脚本一键写入配置："
    )
    _consolas(doc, "python scripts/setup_mcp.py")
    doc.add_heading("一键写入 Agent 配置", level=3)
    doc.add_paragraph(
        "也可以运行交互式脚本："
    )
    _consolas(doc, "python scripts/setup_mcp.py")
    doc.add_paragraph(
        "脚本会自动生成配置，并写入所选 Agent 的配置文件。"
    )

    # 4.8 REST API 接口
    doc.add_heading("4.8 REST API 接口", level=2)
    doc.add_paragraph(
        "提供完整的 RESTful API，支持系统集成。"
        "所有接口需要 JWT 认证（登录获取 Token 后放在 Authorization Header 中）。"
        "支持分页查询、按标签/类型筛选、排序、版本管理、批量导出、异步任务、Wiki 操作等。"
    )

    # 4.9 安全与审计
    doc.add_heading("4.9 安全与审计", level=2)
    for s in [
        "操作日志：所有写操作（create/update/delete/ingest）自动记录，支持查询和筛选",
        "预览模式：写操作前可先 dry_run 预览变更内容，确认无误后再执行",
        "撤销操作：支持撤销 create（软删）、update（恢复字段）、delete（恢复条目）、ingest（恢复）",
        "软删除与恢复：删除的知识条目进入 .trash 目录，可随时通过 restore_knowledge 恢复",
        "API Key 安全：通过系统 Keychain 安全保管，JWT Token 自动生成和管理",
    ]:
        doc.add_paragraph(s, style="List Bullet")

    # ==== 5. 配置说明 ====
    doc.add_heading("5. 配置说明", level=1)
    doc.add_paragraph("配置文件为项目根目录下的 config.yaml，也可通过 GUI 设置界面修改。")

    doc.add_heading("5.1 LLM 配置", level=2)
    _add_table(doc, ["字段", "说明", "示例"], [
        ["provider", "供应商名称（自定义标识）", "deepseek"],
        ["api_key", "API 密钥", "sk-xxx..."],
        ["base_url", "API 地址（必填）", "https://api.deepseek.com/v1"],
        ["model", "模型名称", "deepseek-chat"],
        ["temperature", "创造性程度 0-1", "0.7"],
        ["max_tokens", "最大输出 token 数", "2048"],
        ["timeout", "请求超时秒数", "60"],
    ])
    doc.add_paragraph()
    doc.add_heading("常见供应商配置", level=3)
    _add_table(doc, ["供应商", "API 地址", "模型示例"], [
        ("DeepSeek", "https://api.deepseek.com/v1", "deepseek-chat"),
        ("智谱 GLM", "https://open.bigmodel.cn/api/paas/v4", "glm-4-flash"),
        ("Moonshot", "https://api.moonshot.cn/v1", "moonshot-v1-8k"),
        ("硅基流动", "https://api.siliconflow.cn/v1", "deepseek-ai/DeepSeek-V3"),
        ("MiniMax", "https://api.minimaxi.com/v1", "MiniMax-M3"),
        ("Ollama 本地", "http://localhost:11434/v1", "qwen2"),
    ])

    doc.add_heading("5.2 Embedding 配置", level=2)
    doc.add_paragraph(
        "Embedding 模型用于将文本转换为向量，用于语义搜索。"
        "大多数供应商的 Embedding 接口与 LLM 共享同一地址和 Key，只需修改模型名即可。"
        "在设置界面勾选「与 LLM 使用相同供应商」可自动复用配置。默认使用 BAAI/bge-m3（1024 维）。"
    )

    doc.add_heading("5.3 RAG 管线参数", level=2)
    _add_table(doc, ["参数", "说明", "默认值"], [
        ["top_k", "检索返回的知识块数量", "8"],
        ["chunk_size", "文本分块大小（字符数）", "1200"],
        ["chunk_overlap", "分块重叠字符数", "180"],
        ["score_threshold", "相似度阈值（0-1）", "0.35"],
        ["vector_weight", "向量搜索权重", "0.7"],
        ["keyword_weight", "关键词搜索权重", "0.3"],
        ["enable_rerank", "启用 LLM 重排序", "true"],
        ["enable_query_rewriting", "启用查询改写", "true"],
    ])

    # ==== 6. API 接口参考 ====
    doc.add_heading("6. API 接口参考", level=1)
    doc.add_paragraph("API 文档地址：http://localhost:8000/docs")
    doc.add_paragraph("所有接口需在 Header 中携带 Authorization: Bearer <token>")

    api_routes = [
        ("POST", "/api/auth/register", "注册用户"),
        ("POST", "/api/auth/login", "登录获取 Token"),
        ("GET", "/api/health", "健康检查"),
        ("GET", "/api/knowledge", "知识列表（分页/筛选/排序）"),
        ("GET", "/api/knowledge/search?q=xxx", "搜索知识"),
        ("GET", "/api/knowledge/tags", "获取所有标签"),
        ("GET", "/api/knowledge/{id}", "知识详情"),
        ("POST", "/api/knowledge", "创建知识"),
        ("PUT", "/api/knowledge/{id}", "更新知识"),
        ("DELETE", "/api/knowledge/{id}", "删除知识"),
        ("GET", "/api/knowledge/{id}/versions", "版本历史"),
        ("POST", "/api/knowledge/{id}/versions/{v}/restore", "恢复版本"),
        ("POST", "/api/knowledge/export", "批量导出"),
        ("POST", "/api/chat/ask", "RAG 智能问答"),
        ("GET", "/api/chat/conversations", "对话列表"),
        ("GET", "/api/chat/conversations/{id}/messages", "对话消息"),
    ]
    _add_table(doc, ["方法", "路径", "说明"], api_routes)

    # ==== 7. MCP 工具一览 ====
    doc.add_heading("7. MCP 工具一览", level=1)
    doc.add_paragraph(
        "MCP Server 当前注册 51 个原始工具、51 个命名空间别名、3 个资源和 5 个 Prompt。以下为常用工具分类："
    )

    doc.add_heading("核心知识操作", level=2)
    _add_table(doc, ["工具", "说明"], [
        ("search", "语义向量搜索知识库（Wiki 优先）"),
        ("search_fulltext", "FTS5 全文关键词搜索"),
        ("ask", "RAG 智能问答（6 阶段管线）"),
        ("create", "创建知识条目（自动分块、向量化索引）"),
        ("read", "读取知识条目完整信息"),
        ("update", "更新知识条目（自动版本快照）"),
        ("delete", "软删除知识条目"),
        ("list_knowledge", "列出知识条目（筛选/分页/排序）"),
        ("tags", "获取所有标签"),
        ("restore_knowledge", "恢复已软删除的条目"),
        ("reindex_all", "重建全部索引"),
    ])

    doc.add_heading("文件导入与异步任务", level=2)
    _add_table(doc, ["工具", "说明"], [
        ("ingest_file", "导入本地文件（PDF/DOCX/XLSX/PPTX 等）"),
        ("ingest_url", "导入网页内容"),
        ("create_ingest_job", "创建异步导入任务"),
        ("get_job / list_jobs / cancel_job", "异步任务管理"),
    ])

    doc.add_heading("Wiki 系统", level=2)
    _add_table(doc, ["工具", "说明"], [
        ("save_to_wiki", "将问答结果保存为 Wiki 页面"),
        ("wiki_lint", "Wiki 健康检查"),
        ("wiki_submit_review / approve / reject / deprecate", "Wiki 工作流管理"),
        ("wiki_list_versions / restore_version", "Wiki 版本管理"),
    ])

    doc.add_heading("查询与图谱", level=2)
    _add_table(doc, ["工具", "说明"], [
        ("route_query", "路由分析（推荐执行模式）"),
        ("execute_query", "执行 QuerySpec DSL"),
        ("ask_with_query", "用 DSL 控制 RAG 检索阶段"),
        ("get_source_graph", "构建 RAG 证据链图谱"),
        ("structured_query", "结构化条件过滤"),
        ("explain_query", "解释查询执行计划"),
        ("graph_traverse", "从节点遍历知识图谱"),
    ])

    doc.add_heading("安全与审计", level=2)
    _add_table(doc, ["工具", "说明"], [
        ("kb_capabilities", "查询 MCP 能力清单"),
        ("preview_operation", "预览写操作（不执行）"),
        ("query_operation_logs", "查询操作审计日志"),
        ("undo_operation", "撤销操作"),
        ("get_operation_log", "查询单条操作日志详情"),
    ])

    doc.add_heading("MCP 资源", level=2)
    _add_table(doc, ["URI", "说明"], [
        ("kb://knowledge/{id}", "获取指定知识条目完整内容"),
        ("kb://tags", "获取所有标签"),
        ("kb://stats", "获取知识库统计信息"),
    ])

    # ==== 8. 版本更新日志 ====
    doc.add_heading("8. 版本更新日志", level=1)

    doc.add_heading(f"v{VERSION}（{datetime.now().strftime('%Y-%m-%d')}）", level=2)
    for c in [
        "Wiki 知识沉淀系统（ingest/merge/link/save）",
        "知识图谱可视化（6 种关系类型，LLM 自动分类 15 大类 57 小类）",
        "Block 级引用体系（段落级嵌入 Transclusion）",
        "DSL 结构化查询语言（structured/graph/hybrid 三种模式）",
        "异步任务系统（大文件后台导入、进度追踪）",
        "Reranker 重排序支持（LLM 打分 + 硅基流动模型）",
        "MCP HTTP 常驻服务模式（scripts/mcp_service.py）",
        "REST API 接口扩展（Wiki / 异步任务 / 实时查询）",
        "安全审计闭环（操作日志 / 预览模式 / 撤销操作）",
        "DI 容器架构重构（AppContainer 统一依赖注入）",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("v1.0.0（2026-04-18）", level=2)
    for c in [
        "初始发布，包含核心功能：知识库 CRUD、混合搜索（向量 + 关键词 + RRF 融合）、RAG 智能问答",
        "多格式文件导入（PDF/DOCX/XLSX/HTML/图片）",
        "PySide6 桌面客户端、FastAPI REST API、FastMCP Server",
        "SQLite + FTS5 存储、JWT 认证、keyring 安全密钥管理",
        "Windows 安装包（Inno Setup）、Docker 容器化部署",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    # ==== 9. 常见问题 ====
    doc.add_heading("9. 常见问题", level=1)
    faqs = [
        ("Q: 首次启动提示缺少依赖？",
         "A: 请确保已安装 Python 3.10+，并运行 pip install -e \".[all]\" 安装全部依赖。"),
        ("Q: 导入 PDF 文件时中文乱码？",
         "A: 系统使用 PyPDF2 解析 PDF，部分扫描版 PDF 可能无法正确提取文字。建议使用文字版 PDF。"),
        ("Q: API Key 填写后仍然报错？",
         "A: 请检查 API 地址是否正确（需要以 /v1 结尾），以及 API Key 是否有效。可在设置界面测试。"),
        ("Q: Docker 部署如何持久化数据？",
         "A: docker-compose.yml 已将 ./data 目录挂载到容器内，数据会保存在宿主机。"),
        ("Q: 如何切换 LLM 供应商？",
         "A: 在设置界面修改供应商名称、API 地址、Key 和模型名即可，所有供应商统一使用 OpenAI 兼容协议。"),
        ("Q: MCP HTTP 模式如何常驻运行？",
         "A: 运行 python scripts/mcp_service.py start 即可后台启动，关闭 GUI/终端不影响服务。"
         "用 status 查看状态、stop 停止、restart 重启。"),
        ("Q: 如何自定义 RAG 管线阶段？",
         "A: 在 config.yaml 的 rag.pipeline.stages 中配置阶段列表，也可通过 rag.custom_stages 添加自定义阶段。"),
        ("Q: Wiki 编译消耗太多 LLM 调用怎么办？",
         "A: 通过 wiki.max_llm_calls_per_ingest 限制每次导入的 LLM 调用次数。"
         "也可使用 wiki_lint 工具进行健康检查，手动修复问题页面。"),
    ]
    for q, a in faqs:
        doc.add_paragraph(q).runs[0].font.bold = True
        doc.add_paragraph(a)

    # ---- 保存 ----
    os.makedirs("docs", exist_ok=True)
    filename = f"docs/{APP_NAME}_UserManual_v{VERSION}.docx"
    doc.save(filename)
    print(f"文档已生成: {filename}")
    return filename


if __name__ == "__main__":
    build_doc()
